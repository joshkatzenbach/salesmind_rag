"""
File processing module for converting various file types to plain text.
Supports PDF, DOC, DOCX, and TXT files.
"""

import io
import logging
from typing import Optional
from fastapi import UploadFile, HTTPException

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class FileProcessor:
    """Handles conversion of various file types to plain text."""
    
    @staticmethod
    async def process_file(file: UploadFile) -> str:
        """
        Process uploaded file and convert to plain text.
        
        Args:
            file: Uploaded file from FastAPI
            
        Returns:
            str: Extracted plain text content
            
        Raises:
            HTTPException: If file type is not supported or processing fails
        """
        if not file.filename:
            raise HTTPException(status_code=400, detail="No filename provided")
        
        # Get file extension
        file_extension = file.filename.lower().split('.')[-1]
        
        # Read file content
        file_content = await file.read()
        
        try:
            if file_extension == 'txt':
                return await FileProcessor._process_txt(file_content)
            elif file_extension == 'pdf':
                return await FileProcessor._process_pdf(file_content)
            elif file_extension in ['doc', 'docx']:
                return await FileProcessor._process_doc(file_content, file_extension)
            else:
                raise HTTPException(
                    status_code=400, 
                    detail=f"Unsupported file type: {file_extension}. Supported types: txt, pdf, doc, docx"
                )
        except Exception as e:
            logger.error(f"Error processing file {file.filename}: {str(e)}")
            raise HTTPException(status_code=500, detail=f"Error processing file: {str(e)}")
    
    @staticmethod
    async def _process_txt(file_content: bytes) -> str:
        """Process TXT files."""
        try:
            return file_content.decode('utf-8')
        except UnicodeDecodeError:
            # Try with different encodings
            for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    return file_content.decode(encoding)
                except UnicodeDecodeError:
                    continue
            raise HTTPException(status_code=400, detail="Unable to decode text file")
    
    @staticmethod
    async def _process_pdf(file_content: bytes) -> str:
        """Process PDF files using PyPDF2."""
        try:
            import PyPDF2
        except ImportError:
            raise HTTPException(
                status_code=500, 
                detail="PyPDF2 is required for PDF processing. Install with: pip install PyPDF2"
            )
        
        try:
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
            text = ""
            
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text += page.extract_text() + "\n"
            
            if not text.strip():
                raise HTTPException(status_code=400, detail="No text content found in PDF")
            
            return text.strip()
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Error reading PDF: {str(e)}")
    
    @staticmethod
    async def _process_doc(file_content: bytes, file_extension: str) -> str:
        """Process DOC and DOCX files using python-docx."""
        try:
            from docx import Document
        except ImportError:
            raise HTTPException(
                status_code=500, 
                detail="python-docx is required for DOC/DOCX processing. Install with: pip install python-docx"
            )
        
        try:
            if file_extension == 'docx':
                # Process DOCX files
                doc = Document(io.BytesIO(file_content))
                text = ""
                
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"
                
                if not text.strip():
                    raise HTTPException(status_code=400, detail="No text content found in DOCX")
                
                return text.strip()
            
            elif file_extension == 'doc':
                # For DOC files, we'll try to read as DOCX first, then fallback
                try:
                    doc = Document(io.BytesIO(file_content))
                    text = ""
                    
                    for paragraph in doc.paragraphs:
                        text += paragraph.text + "\n"
                    
                    if text.strip():
                        return text.strip()
                except:
                    pass
                
                # If DOCX reading fails, try to read as plain text
                try:
                    return file_content.decode('utf-8', errors='ignore')
                except:
                    raise HTTPException(status_code=400, detail="Unable to process DOC file")
            
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Error reading {file_extension.upper()} file: {str(e)}")
    
    @staticmethod
    def get_supported_extensions() -> list:
        """Get list of supported file extensions."""
        return ['txt', 'pdf', 'doc', 'docx']
    
    @staticmethod
    def is_supported_file(file: UploadFile) -> bool:
        """Check if the uploaded file type is supported."""
        if not file.filename:
            return False
        
        file_extension = file.filename.lower().split('.')[-1]
        return file_extension in FileProcessor.get_supported_extensions()
